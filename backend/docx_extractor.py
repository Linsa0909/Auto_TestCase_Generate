"""
Word Document Extractor
Extracts text and embedded images from .docx files.
Text: extracted from paragraphs (with heading awareness) and tables.
Images: large screenshots (>10KB, >100x100px) are OCR'd.
Small icons/badges are skipped.
"""

import logging
import io
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageEnhance, ImageFilter

logger = logging.getLogger(__name__)

# Minimums for an image to be worth OCR'ing
MIN_IMAGE_BYTES = 8192       # 8KB — skip icons & tiny badges
MIN_IMAGE_WIDTH = 100
MIN_IMAGE_HEIGHT = 100


def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """Lightweight preprocessing to improve OCR accuracy."""
    # Convert to grayscale
    if img.mode != "L":
        img = img.convert("L")
    # Enhance contrast
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.8)
    # Sharpen lightly
    img = img.filter(ImageFilter.SHARPEN)
    return img


async def extract_docx(file_path: str) -> str:
    """
    Extract all content from a .docx file:
    - Structured text from paragraphs (grouped by heading level)
    - Table data
    - OCR text from large embedded images (screenshots)
    Returns formatted semantic text ready for AI generation.
    """
    parts = []

    # --- 1. Extract structured text ---
    try:
        doc = Document(file_path)

        # Build structured output with heading awareness
        para_groups = []
        current_group = {"heading": "", "lines": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Detect headings by style or short non-punctuated lines
            is_heading = (
                style_name.startswith("Heading") or
                style_name == "Title" or
                (len(text) <= 20 and not any(p in text for p in ("，", "。", "：", "；", "、")) and para.runs and para.runs[0].bold)
            )

            if is_heading and current_group["lines"]:
                para_groups.append(current_group)
                current_group = {"heading": text, "lines": []}
            elif is_heading:
                if current_group["heading"]:
                    para_groups.append(current_group)
                current_group = {"heading": text, "lines": []}
            else:
                current_group["lines"].append(text)

        if current_group["heading"] or current_group["lines"]:
            para_groups.append(current_group)

        # Format groups
        if para_groups:
            parts.append("=== 需求文档结构化内容 ===")
            for g in para_groups:
                if g["heading"]:
                    parts.append(f"\n## {g['heading']}")
                parts.append("\n".join(g["lines"]))

        # Tables (keep existing table logic)
        table_texts = []
        for ti, table in enumerate(doc.tables):
            table_lines = [f"\n[表格 {ti + 1}]"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            table_texts.append("\n".join(table_lines))

        if table_texts:
            parts.append("\n\n=== 需求文档表格内容 ===")
            parts.extend(table_texts)

        para_count = sum(len(g["lines"]) for g in para_groups)
        logger.info(f"Extracted {para_count} paragraphs ({len(para_groups)} groups), {len(doc.tables)} tables from {file_path}")

    except Exception as e:
        logger.error(f"Failed to extract text from docx: {e}")
        parts.append("(Word 文档文字提取失败)")

    # --- 2. Extract large images and OCR ---
    try:
        from ocr_extractor import extract_text as ocr_text

        ocr_parts = []
        skipped_small = 0
        skipped_dim = 0
        image_count = 0

        with ZipFile(file_path, "r") as zf:
            for name in zf.namelist():
                if not (name.startswith("word/media/") and any(
                    name.lower().endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')
                )):
                    continue

                info = zf.getinfo(name)
                image_bytes = zf.read(name)

                # Skip tiny images (icons, badges, bullet points)
                if info.file_size < MIN_IMAGE_BYTES:
                    skipped_small += 1
                    continue

                try:
                    img = Image.open(io.BytesIO(image_bytes))
                    w, h = img.size
                    if w < MIN_IMAGE_WIDTH or h < MIN_IMAGE_HEIGHT:
                        skipped_dim += 1
                        continue

                    # Preprocess for better OCR
                    img = _preprocess_for_ocr(img)
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    png_bytes = buf.getvalue()

                    result = await ocr_text(png_bytes)
                    if result and "未识别" not in result:
                        # Clean up OCR noise — keep if any meaningful content found
                        lines = result.split("\n")
                        meaningful = [l for l in lines if len(l.strip()) > 2 and not l.startswith("[第")]
                        if len(meaningful) >= 1:
                            ocr_parts.append(result)
                            image_count += 1
                        else:
                            skipped_small += 1
                except Exception as e:
                    logger.warning(f"OCR failed for docx image {name}: {e}")

        if image_count > 0:
            parts.append(f"\n\n=== 文档截图OCR识别结果 ({image_count}张) ===")
            parts.extend(ocr_parts)

        logger.info(
            f"OCR: {image_count} images, skipped {skipped_small} small + {skipped_dim} tiny-dimension "
            f"from {file_path}"
        )

    except Exception as e:
        logger.error(f"Failed to extract images from docx: {e}")

    if not parts:
        return "(Word 文档未提取到有效内容)"

    return "\n".join(parts)
